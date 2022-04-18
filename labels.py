import time
from abc import ABCMeta, abstractmethod
from typing import Any, List

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import image_size
from docx_helper import *


class Label(metaclass=ABCMeta):
    """
    内容标签接口
    """

    @classmethod
    @abstractmethod
    def get_type(cls) -> str:
        """获取内容标签类型"""
        pass

    @classmethod
    @abstractmethod
    def has_content(cls) -> bool:
        """是否有插入内容（生成时外部输入的信息）"""
        pass

    @classmethod
    @abstractmethod
    def register_static_datas(cls, static_datas: dict) -> None:
        """注册静态数据"""
        pass

    @classmethod
    @abstractmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        """在插入点插入数据"""
        pass

    @classmethod
    @abstractmethod
    def check_data_type(cls, data: Any) -> bool:
        """检查插入输入类型"""
        pass


class LabelManager:
    __labels = []

    @classmethod
    def register(cls, label):
        if not issubclass(label, Label):
            raise TypeError('label must be sub class of type `Label`')
        cls.__labels.append(label)

    @classmethod
    def get_labels(cls) -> List[Label]:
        return cls.__labels.copy()

    @classmethod
    def print_registered_labels(cls):
        print(f'当前注册的内容标签类型有：{[l.get_type() for l in cls.__labels]}')
        print(f'无内容类型有：{[l.get_type() for l in cls.__labels if not l.has_content()]}')
        print(f'有内容类型有：{[l.get_type() for l in cls.__labels if l.has_content()]}')


class TextLabel(Label):
    """文本内容标签"""

    @classmethod
    def get_type(cls) -> str:
        return 'text'

    @classmethod
    def has_content(cls) -> bool:
        return True

    @classmethod
    def register_static_datas(cls, static_datas: dict) -> None:
        pass

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        new_text = point_data['run'].text.replace(point_data['text'], data)
        point_data['run'].text = new_text

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        return isinstance(data, str)


LabelManager.register(TextLabel)


class DateLabel(Label):
    @classmethod
    def get_type(cls) -> str:
        return 'date'

    @classmethod
    def has_content(cls) -> bool:
        return False

    @classmethod
    def register_static_datas(cls, static_datas: dict) -> None:
        now_time = time.time()
        now_time_struct = time.localtime(now_time)
        date_s = time.strftime("%Y-%m-%d", now_time_struct)
        static_datas[cls.get_type()] = date_s

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        new_text = point_data['run'].text.replace(point_data['text'], static_datas[point_data['type']])
        point_data['run'].text = new_text

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        return True


LabelManager.register(DateLabel)


class TimeLabel(Label):
    @classmethod
    def get_type(cls) -> str:
        return 'time'

    @classmethod
    def has_content(cls) -> bool:
        return False

    @classmethod
    def register_static_datas(cls, static_datas: dict) -> None:
        now_time = time.time()
        now_time_struct = time.localtime(now_time)
        time_s = time.strftime("%H:%M:%S", now_time_struct)
        static_datas[cls.get_type()] = time_s

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        new_text = point_data['run'].text.replace(point_data['text'], static_datas[point_data['type']])
        point_data['run'].text = new_text

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        return True


LabelManager.register(TimeLabel)


class OrderedListLabel(Label):
    @classmethod
    def get_type(cls) -> str:
        return 'ordered-list'

    @classmethod
    def has_content(cls) -> bool:
        return True

    @classmethod
    def register_static_datas(cls, static_datas: dict) -> None:
        pass

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        paragraph = point_data['paragraph']
        for i, item in enumerate(data):
            p = paragraph.insert_paragraph_before(f'{i + 1}. {item}')
            copy_paragraph_style(p, paragraph)
        delete_paragraph(paragraph)

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        """要求data是list，且元素是str"""
        return isinstance(data, list) and all([isinstance(d, str) for d in data])


LabelManager.register(OrderedListLabel)


class UnorderedListLabel(Label):
    _header_chars = {"circle0": "•", "square0": "▪", "disc0": "◦",
                     "circle1": "●", "square1": "■", "disc1": "○", "diamond1": "◆", "diamond1_e": "◇", }
    _default_header_char = "circle1"
    _default_header_gap = 1

    @classmethod
    def get_type(cls) -> str:
        return 'unordered-list'

    @classmethod
    def has_content(cls) -> bool:
        return True

    @classmethod
    def register_static_datas(cls, static_datas: dict) -> None:
        pass

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        paragraph = point_data['paragraph']
        for i, item in enumerate(data):
            p = paragraph.insert_paragraph_before(f'{cls._header_chars[cls._default_header_char]}{" " * cls._default_header_gap}{item}')
            copy_paragraph_style(p, paragraph)
        delete_paragraph(paragraph)

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        """要求data是list，且元素是str"""
        return isinstance(data, list) and all([isinstance(d, str) for d in data])


LabelManager.register(UnorderedListLabel)


class ImageLabel(Label):
    @classmethod
    def get_type(cls) -> str:
        return 'image'

    @classmethod
    def has_content(cls) -> bool:
        return True

    @classmethod
    def register_static_datas(cls, static_datas: dict) -> None:
        pass

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        """当前仅支持url为本地路径的图片，且图片单独成段"""
        pic_desc, pic_url = data
        document = point_data['document']
        paragraph = point_data['paragraph']
        run = point_data['run']

        # 默认某方向完全填充文档自动调整图片大小
        # 默认使用第一节的宽高当作整个文档的宽高
        d_section = document.sections[0]
        doc_height, doc_width = d_section.page_height, d_section.page_width
        l_margin, r_margin, t_margin, b_margin = d_section.left_margin, d_section.right_margin, d_section.top_margin, d_section.bottom_margin
        max_height = doc_height - t_margin - b_margin
        max_width = doc_width - l_margin - r_margin
        img_width, img_height = image_size.get(pic_url)

        limit_doc_width = True
        if doc_width / doc_height > img_width / img_height:
            limit_doc_width = False

        ip = paragraph.insert_paragraph_before()
        ir = ip.add_run()
        if limit_doc_width:
            ir.add_picture(pic_url, width=max_width)
        else:
            ir.add_picture(pic_url, height=max_height)

        if pic_desc is None:
            delete_paragraph(paragraph)
        else:
            run.text = pic_desc
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        """要求data为tuple，第一个元素是描述字符串，第二个元素是图片url"""
        return isinstance(data, tuple) and len(data) == 2 and (isinstance(data[0], str) or data[0] is None) and isinstance(data[1], str)


LabelManager.register(ImageLabel)
